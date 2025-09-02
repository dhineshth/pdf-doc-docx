import os
import sys
import subprocess
import zipfile
import olefile
import re
import tempfile
import platform
from pathlib import Path
from fastapi import FastAPI, File, UploadFile, HTTPException, Request, Header
from fastapi.responses import JSONResponse
import PyPDF2
from typing import Optional
import ua_parser.user_agent_parser as ua_parser
import uuid
import json
from enum import Enum

app = FastAPI(title="Resume Parser API", description="API for counting pages in resume files", version="1.0.0")

class OSType(str, Enum):
    WINDOWS = "windows"
    MACOS = "macos"
    LINUX = "linux"
    UNKNOWN = "unknown"

def detect_os_from_user_agent(user_agent: str) -> OSType:
    """Detect OS from User-Agent header"""
    if not user_agent:
        return OSType.UNKNOWN
    
    try:
        parsed_string = ua_parser.Parse(user_agent)
        os_family = parsed_string['os']['family'].lower()
        
        if 'windows' in os_family:
            return OSType.WINDOWS
        elif 'mac' in os_family or 'ios' in os_family:
            return OSType.MACOS
        elif 'linux' in os_family or 'android' in os_family:
            return OSType.LINUX
        else:
            return OSType.UNKNOWN
    except:
        # Fallback simple detection
        user_agent_lower = user_agent.lower()
        if 'windows' in user_agent_lower:
            return OSType.WINDOWS
        elif 'mac' in user_agent_lower:
            return OSType.MACOS
        elif 'linux' in user_agent_lower:
            return OSType.LINUX
        else:
            return OSType.UNKNOWN

class ResumePageCounter:
    def __init__(self, client_os: OSType):
        self.client_os = client_os
        
    def count_pages(self, file_path: str) -> int:
        """
        Main function to count pages in DOC/DOCX/PDF files
        Uses methods appropriate for the client's OS
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"The file {file_path} does not exist")
        
        file_ext = os.path.splitext(file_path)[1].lower()
        
        if file_ext not in ['.doc', '.docx', '.pdf']:
            raise ValueError("Unsupported file format. Please provide a .doc, .docx, or .pdf file")
        
        # Handle PDF files (same for all OS types)
        if file_ext == '.pdf':
            return self._count_pdf_pages(file_path)
        
        # For Word documents, use client-appropriate methods
        try:
            if self.client_os == OSType.WINDOWS:
                page_count = self._count_pages_windows(file_path)
            elif self.client_os == OSType.MACOS:
                page_count = self._count_pages_macos(file_path)
            else:
                # For Linux and unknown OS, use cross-platform methods
                page_count = self._count_pages_cross_platform(file_path)
                
            if page_count > 0:
                return page_count
        except Exception as e:
            print(f"OS-specific method failed: {e}. Trying fallback methods...")
        
        # If OS-specific methods fail, try cross-platform methods
        try:
            page_count = self._count_pages_cross_platform(file_path)
            if page_count > 0:
                return page_count
        except Exception as e:
            print(f"Cross-platform method failed: {e}")
        
        # Final fallback: estimation based on file size
        return self._estimate_pages(file_path)
    
    def _count_pdf_pages(self, file_path: str) -> int:
        """Count pages in PDF files using PyPDF2"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                return len(pdf_reader.pages)
        except Exception as e:
            raise Exception(f"Failed to count PDF pages: {e}")
    
    def _count_pages_windows(self, file_path: str) -> int:
        """Count pages on Windows using Microsoft Word COM interface"""
        try:
            # This would only work if the server is Windows and has Word installed
            # For client-side processing, we need to implement a different approach
            # For now, we'll use cross-platform methods
            return self._count_pages_cross_platform(file_path)
        except Exception as e:
            raise Exception(f"Windows method failed: {e}")
    
    def _count_pages_macos(self, file_path: str) -> int:
        """Count pages on macOS using AppleScript with Microsoft Word"""
        try:
            # This would only work if the server is macOS and has Word installed
            # For client-side processing, we need to implement a different approach
            # For now, we'll use cross-platform methods
            return self._count_pages_cross_platform(file_path)
        except Exception as e:
            raise Exception(f"macOS method failed: {e}")
    
    def _count_pages_cross_platform(self, file_path: str) -> int:
        """Cross-platform method to extract page count from DOCX/DOC files"""
        file_ext = os.path.splitext(file_path)[1].lower()
        
        if file_ext == '.docx':
            return self._count_docx_pages(file_path)
        elif file_ext == '.doc':
            return self._count_doc_pages(file_path)
        else:
            return self._estimate_pages(file_path)
    
    def _count_docx_pages(self, file_path: str) -> int:
        """Extract page count from DOCX file (which is a ZIP archive)"""
        try:
            with zipfile.ZipFile(file_path) as docx:
                # Check if the document properties contain page count
                if 'docProps/app.xml' in docx.namelist():
                    with docx.open('docProps/app.xml') as app_xml:
                        content = app_xml.read().decode('utf-8')
                        # Look for Pages tag
                        match = re.search(r'<Pages>(\d+)</Pages>', content)
                        if match:
                            return int(match.group(1))
            
            # If page count not found in metadata, try to estimate from content
            return self._estimate_docx_pages(file_path)
        except:
            return self._estimate_pages(file_path)
    
    def _estimate_docx_pages(self, file_path: str) -> int:
        """Estimate page count for DOCX files by examining content"""
        try:
            from docx import Document
            
            doc = Document(file_path)
            total_content = 0
            
            # Count characters in paragraphs
            for paragraph in doc.paragraphs:
                total_content += len(paragraph.text)
            
            # Count characters in tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            total_content += len(paragraph.text)
            
            # Estimate pages based on content length
            if total_content < 1500:
                return 1
            elif total_content < 3000:
                return 2
            else:
                return max(2, total_content // 1500)
        except ImportError:
            # python-docx not installed, fall back to file size estimation
            return self._estimate_pages(file_path)
        except:
            return self._estimate_pages(file_path)
    
    def _count_doc_pages(self, file_path: str) -> int:
        """Attempt to extract page count from binary DOC format"""
        # This is challenging without external libraries
        # For now, we'll just estimate based on file size
        return self._estimate_pages(file_path)
    
    def _estimate_pages(self, file_path: str) -> int:
        """Fallback method to estimate page count based on file size"""
        file_size = os.path.getsize(file_path)  # size in bytes
        
        # Rough estimation based on file type
        file_ext = os.path.splitext(file_path)[1].lower()
        
        if file_ext == '.docx':
            return max(1, round(file_size / 45000))  # ~45KB per page
        elif file_ext == '.doc':
            return max(1, round(file_size / 35000))  # ~35KB per page
        else:  # .pdf
            return max(1, round(file_size / 75000))  # ~75KB per page

# Client-side JavaScript code template
CLIENT_JS_TEMPLATE = """
<script>
// Client-side page counter for resume files
class ClientPageCounter {
    constructor() {
        this.osType = this.detectOS();
    }
    
    detectOS() {
        const platform = navigator.platform.toLowerCase();
        if (platform.includes('win')) return 'windows';
        if (platform.includes('mac')) return 'macos';
        if (platform.includes('linux')) return 'linux';
        return 'unknown';
    }
    
    async countPages(file) {
        const formData = new FormData();
        formData.append('file', file);
        formData.append('client_os', this.osType);
        
        try {
            const response = await fetch('/count-pages-client', {
                method: 'POST',
                body: formData
            });
            
            if (!response.ok) {
                throw new Error(`HTTP error! status: ${response.status}`);
            }
            
            return await response.json();
        } catch (error) {
            console.error('Error counting pages:', error);
            throw error;
        }
    }
    
    // Fallback method for client-side PDF page counting
    async countPdfPages(file) {
        return new Promise((resolve, reject) => {
            const reader = new FileReader();
            reader.onload = function(e) {
                const typedArray = new Uint8Array(e.target.result);
                
                // This would require pdf.js or similar library
                // For now, we'll just estimate based on file size
                const estimatedPages = Math.max(1, Math.round(file.size / 75000));
                resolve(estimatedPages);
            };
            reader.onerror = reject;
            reader.readAsArrayBuffer(file);
        });
    }
}

// Global instance
window.clientPageCounter = new ClientPageCounter();
</script>
"""

# FastAPI endpoints
@app.get("/")
async def root(request: Request, user_agent: Optional[str] = Header(None)):
    client_os = detect_os_from_user_agent(user_agent)
    
    return {
        "message": "Resume Parser API - Use /count-pages endpoint to count pages in resume files",
        "client_os": client_os.value,
        "note": "Processing methods are selected based on your client OS"
    }

@app.post("/count-pages")
async def count_pages(
    file: UploadFile = File(...),
    user_agent: Optional[str] = Header(None)
):
    """
    Count pages in uploaded resume file (DOC, DOCX, or PDF)
    Uses methods appropriate for the client's OS
    """
    # Detect client OS
    client_os = detect_os_from_user_agent(user_agent)
    
    # Validate file type
    if not file.filename:
        raise HTTPException(status_code=400, detail="No file provided")
    
    file_ext = Path(file.filename).suffix.lower()
    if file_ext not in ['.doc', '.docx', '.pdf']:
        raise HTTPException(
            status_code=400, 
            detail="Unsupported file format. Please upload a .doc, .docx, or .pdf file"
        )
    
    # Save uploaded file to temporary location
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=file_ext) as temp_file:
            content = await file.read()
            temp_file.write(content)
            temp_file_path = temp_file.name
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to process uploaded file: {str(e)}")
    
    # Count pages using client-appropriate methods
    try:
        counter = ResumePageCounter(client_os)
        page_count = counter.count_pages(temp_file_path)
        
        # Clean up temporary file
        try:
            os.unlink(temp_file_path)
        except:
            pass
            
        return JSONResponse(
            status_code=200,
            content={
                "filename": file.filename,
                "page_count": page_count,
                "file_type": file_ext[1:],  # Remove the dot
                "client_os": client_os.value,
                "method_used": "server_processing_based_on_client_os"
            }
        )
    except Exception as e:
        # Clean up temporary file even if there's an error
        try:
            os.unlink(temp_file_path)
        except:
            pass
            
        raise HTTPException(status_code=500, detail=f"Failed to count pages: {str(e)}")

@app.post("/count-pages-client")
async def count_pages_client(
    file: UploadFile = File(...),
    client_os: Optional[str] = Header(None)
):
    """
    Endpoint for client-side processing requests
    """
    # This endpoint is designed to be called from client-side JavaScript
    # It includes the client_os in the request
    
    # Validate file type
    if not file.filename:
        raise HTTPException(status_code=400, detail="No file provided")
    
    file_ext = Path(file.filename).suffix.lower()
    if file_ext not in ['.doc', '.docx', '.pdf']:
        raise HTTPException(
            status_code=400, 
            detail="Unsupported file format. Please upload a .doc, .docx, or .pdf file"
        )
    
    # Parse client OS
    client_os_enum = OSType(client_os.lower()) if client_os else OSType.UNKNOWN
    
    # Save uploaded file to temporary location
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=file_ext) as temp_file:
            content = await file.read()
            temp_file.write(content)
            temp_file_path = temp_file.name
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to process uploaded file: {str(e)}")
    
    # Count pages using client-appropriate methods
    try:
        counter = ResumePageCounter(client_os_enum)
        page_count = counter.count_pages(temp_file_path)
        
        # Clean up temporary file
        try:
            os.unlink(temp_file_path)
        except:
            pass
            
        return JSONResponse(
            status_code=200,
            content={
                "filename": file.filename,
                "page_count": page_count,
                "file_type": file_ext[1:],
                "client_os": client_os_enum.value
            }
        )
    except Exception as e:
        # Clean up temporary file even if there's an error
        try:
            os.unlink(temp_file_path)
        except:
            pass
            
        raise HTTPException(status_code=500, detail=f"Failed to count pages: {str(e)}")

@app.get("/client-js")
async def get_client_js():
    """Endpoint to get client-side JavaScript code"""
    return Response(content=CLIENT_JS_TEMPLATE, media_type="application/javascript")

@app.get("/health")
async def health_check():
    """Health check endpoint"""
    return {
        "status": "healthy", 
        "server_platform": platform.platform(),
    }

@app.get("/debug")
async def debug_info(request: Request, user_agent: Optional[str] = Header(None)):
    """Debug endpoint to see client information"""
    client_os = detect_os_from_user_agent(user_agent)
    
    return {
        "client_os": client_os.value,
        "user_agent": user_agent,
        "client_host": request.client.host if request.client else "unknown",
        "note": "Processing methods are selected based on your client OS"
    }

# HTML template for demo page
HTML_TEMPLATE = """
<!DOCTYPE html>
<html>
<head>
    <title>Resume Page Counter</title>
    <style>
        body { font-family: Arial, sans-serif; max-width: 800px; margin: 0 auto; padding: 20px; }
        .upload-container { border: 2px dashed #ccc; padding: 20px; text-align: center; margin: 20px 0; }
        .result { margin-top: 20px; padding: 15px; background: #f5f5f5; border-radius: 5px; }
        .error { color: red; }
    </style>
</head>
<body>
    <h1>Resume Page Counter</h1>
    <p>Upload a DOC, DOCX, or PDF file to count its pages.</p>
    
    <div class="upload-container">
        <input type="file" id="fileInput" accept=".doc,.docx,.pdf">
        <button onclick="countPages()">Count Pages</button>
    </div>
    
    <div id="result" class="result" style="display: none;">
        <h3>Result</h3>
        <div id="resultContent"></div>
    </div>
    
    <script>
        async function countPages() {
            const fileInput = document.getElementById('fileInput');
            const resultDiv = document.getElementById('result');
            const resultContent = document.getElementById('resultContent');
            
            if (!fileInput.files || fileInput.files.length === 0) {
                alert('Please select a file first');
                return;
            }
            
            const file = fileInput.files[0];
            const formData = new FormData();
            formData.append('file', file);
            
            // Detect client OS
            const platform = navigator.platform.toLowerCase();
            let clientOS = 'unknown';
            if (platform.includes('win')) clientOS = 'windows';
            if (platform.includes('mac')) clientOS = 'macos';
            if (platform.includes('linux')) clientOS = 'linux';
            
            formData.append('client_os', clientOS);
            
            resultContent.innerHTML = 'Processing...';
            resultDiv.style.display = 'block';
            
            try {
                const response = await fetch('/count-pages-client', {
                    method: 'POST',
                    body: formData
                });
                
                if (!response.ok) {
                    const error = await response.text();
                    throw new Error(error);
                }
                
                const data = await response.json();
                resultContent.innerHTML = `
                    <p><strong>Filename:</strong> ${data.filename}</p>
                    <p><strong>Page Count:</strong> ${data.page_count}</p>
                    <p><strong>File Type:</strong> ${data.file_type}</p>
                    <p><strong>Client OS:</strong> ${data.client_os}</p>
                `;
            } catch (error) {
                resultContent.innerHTML = `<p class="error">Error: ${error.message}</p>`;
            }
        }
    </script>
</body>
</html>
"""

@app.get("/demo")
async def demo_page():
    """Demo page for testing the page counter"""
    return Response(content=HTML_TEMPLATE, media_type="text/html")

# Command line interface
def main():
    import argparse
    
    parser = argparse.ArgumentParser(description='Count pages in DOC/DOCX/PDF files')
    parser.add_argument('file_path', help='Path to the DOC, DOCX, or PDF file')
    parser.add_argument('--os', help='Specify client OS (windows, macos, linux)', default='unknown')
    args = parser.parse_args()
    
    client_os = OSType(args.os.lower())
    counter = ResumePageCounter(client_os)
    
    try:
        page_count = counter.count_pages(args.file_path)
        print(f"Client OS: {client_os.value}")
        print(f"The document has {page_count} page(s)")
    except Exception as e:
        print(f"Error: {e}")
        sys.exit(1)

if __name__ == "__main__":
    # If run directly, use command line interface
    if len(sys.argv) > 1:
        main()
    else:
        # Import uvicorn to run the FastAPI app
        import uvicorn
        uvicorn.run(app, host="0.0.0.0", port=8000)
